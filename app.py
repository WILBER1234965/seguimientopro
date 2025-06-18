# app.py
"""Aplicación principal Qt con modo claro/oscuro y barra de menús completa."""

import logging
import sys
from PyQt6.QtWidgets import (
    QApplication, QMainWindow, QTabWidget, QMessageBox, QFileDialog,
    QToolBar, QWidget, QCheckBox, QSizePolicy
)
from PyQt6.QtGui import QAction, QIcon
from PyQt6.QtCore import Qt

import pandas as pd
from fpdf import FPDF
from docx import Document

from database       import Database
from dashboard_tab  import DashboardTab
from items_tab      import ItemsTab
from atajados_tab   import AtajadosTab
from avance_tab     import AvanceTab
from cronograma_tab import CronogramaTab
from summary_tab    import SummaryTab


# -------------------  Hojas de estilo  -------------------
LIGHT_QSS = """
    QWidget         { background: #ffffff; color: #202020; font-family: Segoe UI; font-size: 12pt; }
    QGroupBox       { border: 1px solid #ccc; border-radius: 10px; padding: 10px; }
    QPushButton     { background: #1976D2; color: #fff; border-radius: 6px; padding: 6px 14px; }
    QPushButton:hover { background: #1259a4; }
    QLineEdit       { background: #fafafa; border: 1px solid #aaa; border-radius: 6px; padding: 5px; }
    QTableWidget    { background: #f9f9f9; alternate-background-color: #e8f0fe; border: 1px solid #ccc; }
    QHeaderView::section { background: #d0e8ff; padding: 4px; font-weight: bold; border: 1px solid #ccc; }
"""

DARK_QSS = """
    QWidget         { background: #1e1e1e; color: #e0e0e0; font-family: Segoe UI; font-size: 12pt; }
    QGroupBox       { background: #252525; border: 1px solid #444; border-radius: 10px; padding: 10px; }
    QPushButton     { background: #0d6efd; color: #fff; border-radius: 6px; padding: 6px 14px; }
    QPushButton:hover { background: #1a75ff; }
    QLineEdit       { background: #2a2a2a; border: 1px solid #555; border-radius: 6px; padding: 5px; color: #e0e0e0; }
    QTableWidget    { background: #272727; alternate-background-color: #1f1f1f; color: #e0e0e0; border: 1px solid #444; }
    QHeaderView::section { background: #353535; padding: 4px; font-weight: bold; border: 1px solid #444; color:#e0e0e0; }
"""


# -------------------  Toggle de tema  --------------------
class ThemeToggle(QCheckBox):
    """Interruptor personalizado (sol/luna)."""
    def __init__(self):
        super().__init__()
        self.setFixedSize(60, 28)
        self.setTristate(False)
        self.setChecked(False)  # inicia claro
        self.setStyleSheet("QCheckBox::indicator { width:0; height:0; }")
    def paintEvent(self, e):
        from PyQt6.QtGui import QPainter, QColor, QPen
        from math import cos, sin, pi
        p = QPainter(self); p.setRenderHint(QPainter.RenderHint.Antialiasing)
        r = self.rect()
        # fondo
        p.setBrush(QColor("#505050") if self.isChecked() else QColor("#cccccc"))
        p.setPen(Qt.PenStyle.NoPen); p.drawRoundedRect(r, 14, 14)
        # perilla
        knob_x = r.right() - 26 if self.isChecked() else r.left() + 2
        p.setBrush(QColor("#ffffff")); p.drawEllipse(knob_x, 2, 24, 24)
        # icono
        p.setPen(QPen(QColor("#000000"), 2)); cx, cy = knob_x + 12, 14
        if self.isChecked():  # luna
            p.drawArc(cx-6, cy-6, 12, 12, 30*16, 300*16)
        else:                 # sol
            p.drawEllipse(cx-6, cy-6, 12, 12)
            for ang in range(0, 360, 45):
                x1 = cx + 10*cos(ang*pi/180); y1 = cy - 10*sin(ang*pi/180)
                x2 = cx + 14*cos(ang*pi/180); y2 = cy - 14*sin(ang*pi/180)
                p.drawLine(int(x1), int(y1), int(x2), int(y2))
        p.end()


# -------------------  Ventana principal  -----------------
class MainWindow(QMainWindow):
    def __init__(self):
        super().__init__()
        self.db = Database()
        self.setWindowTitle("Supervisión de Atajados")
        self.setWindowIcon(QIcon("resources/icono.png"))  # verifica la ruta
        self.resize(1200, 800)

        # ---------- Pestañas ----------
        self.tabs = QTabWidget()
        self.dashboard_tab  = DashboardTab(self.db)
        self.items_tab      = ItemsTab(self.db)
        self.atajados_tab   = AtajadosTab(self.db)
        self.avance_tab     = AvanceTab(self.db, save_callback=self.refresh_all)
        self.cronograma_tab = CronogramaTab(self.db)
        self.summary_tab    = SummaryTab(self.db)

        self.tabs.addTab(self.dashboard_tab,  "Inicio")
        self.tabs.addTab(self.items_tab,      "Ítems")
        self.tabs.addTab(self.atajados_tab,   "Atajados")
        self.tabs.addTab(self.avance_tab,     "Seguimiento")
        self.tabs.addTab(self.cronograma_tab, "Cronograma")
        self.tabs.addTab(self.summary_tab,    "Resumen")
        self.setCentralWidget(self.tabs)

        # ---------- Menú ----------
        menubar = self.menuBar()
        archivo  = menubar.addMenu("Archivo")
        datos    = menubar.addMenu("Datos")
        estado   = menubar.addMenu("Estado")
        reportes = menubar.addMenu("Reportes")
        exportar = menubar.addMenu("Exportar")

        archivo.addAction("Nuevo trabajo")
        archivo.addAction("Añadir trabajo")
        archivo.addAction("Guardar trabajo")
        datos.addAction("Ítems").triggered.connect(lambda: self.tabs.setCurrentIndex(1))
        datos.addAction("Atajados").triggered.connect(lambda: self.tabs.setCurrentIndex(2))
        estado.addAction("Cronograma").triggered.connect(lambda: self.tabs.setCurrentIndex(4))
        estado.addAction("Seguimiento").triggered.connect(lambda: self.tabs.setCurrentIndex(3))
        reportes.addAction("Generar reporte").triggered.connect(lambda: self.tabs.setCurrentIndex(5))
        exportar.addAction("A Excel").triggered.connect(self.to_excel)
        exportar.addAction("A PDF").triggered.connect(self.to_pdf)
        exportar.addAction("A Word").triggered.connect(self.to_word)

        # ---------- Toolbar + toggle ----------
        toolbar = QToolBar(); toolbar.setMovable(False)
        self.addToolBar(Qt.ToolBarArea.TopToolBarArea, toolbar)
        spacer = QWidget()
        spacer.setSizePolicy(QSizePolicy.Policy.Expanding, QSizePolicy.Policy.Preferred)
        toolbar.addWidget(spacer)
        self.theme_toggle = ThemeToggle()
        self.theme_toggle.toggled.connect(self.apply_theme)  # ← señal correcta
        toolbar.addWidget(self.theme_toggle)
        self.apply_theme(False)  # inicia claro

    # -------------------  Tema -------------------
    def apply_theme(self, checked: bool):
        """Aplicar QSS global; primero limpia para refrescar widgets."""
        app = QApplication.instance()
        app.setStyleSheet("")                                  # limpiar
        app.setStyleSheet(DARK_QSS if checked else LIGHT_QSS)  # nuevo

    # -------------------  Refresh -------------------
    def refresh_all(self):
        self.dashboard_tab.refresh()
        self.items_tab.refresh()
        self.atajados_tab.refresh()
        self.cronograma_tab.refresh()
        self.summary_tab.refresh()

    # -------------------  Cerrar -------------------
    def closeEvent(self, event):
        self.db.close()
        super().closeEvent(event)

    # -------------------  Exportar ------------------
    def to_excel(self):
        path, _ = QFileDialog.getSaveFileName(self, "Guardar", "", "*.xlsx")
        if not path: return
        try:
            df_items = pd.read_sql("SELECT * FROM items", self.db.conn)
            df_ataj  = pd.read_sql("SELECT * FROM atajados", self.db.conn)
            with pd.ExcelWriter(path) as w:
                df_items.to_excel(w, "Ítems",  index=False)
                df_ataj.to_excel(w,  "Atajados", index=False)
            QMessageBox.information(self, "✔", "Excel generado")
        except Exception as exc:
            logging.exception("Error Excel"); QMessageBox.critical(self, "Error", str(exc))

    def to_pdf(self):
        path, _ = QFileDialog.getSaveFileName(self, "Guardar", "", "*.pdf")
        if not path: return
        try:
            pdf = FPDF(); pdf.add_page(); pdf.set_font("Arial", size=12)
            pdf.cell(0, 10, "Reporte Ítems", ln=1)
            for n, t in self.db.fetchall("SELECT name,total FROM items"):
                pdf.cell(0, 8, f"{n}: {t}", ln=1)
            pdf.add_page(); pdf.cell(0, 10, "Reporte Atajados", ln=1)
            for num, com in self.db.fetchall("SELECT number,comunidad FROM atajados"):
                pdf.cell(0, 8, f"{num} - {com}", ln=1)
            pdf.output(path); QMessageBox.information(self, "✔", "PDF generado")
        except Exception as exc:
            logging.exception("Error PDF"); QMessageBox.critical(self, "Error", str(exc))

    def to_word(self):
        path, _ = QFileDialog.getSaveFileName(self, "Guardar", "", "*.docx")
        if not path: return
        try:
            doc = Document(); doc.add_heading("Reporte Ítems", level=1)
            for n, t in self.db.fetchall("SELECT name,total FROM items"):
                doc.add_paragraph(f"{n}: {t}")
            doc.add_page_break(); doc.add_heading("Reporte Atajados", level=1)
            for num, com in self.db.fetchall("SELECT number,comunidad FROM atajados"):
                doc.add_paragraph(f"{num} - {com}")
            doc.save(path); QMessageBox.information(self, "✔", "Word generado")
        except Exception as exc:
            logging.exception("Error Word"); QMessageBox.critical(self, "Error", str(exc))


# -------------------  Lanzador -------------------
if __name__ == "__main__":
    logging.basicConfig(level=logging.INFO, filename="app.log",
                        format="%(asctime)s %(levelname)s %(message)s")
    app = QApplication(sys.argv)
    app.setApplicationName("SupervisiónAtajados")
    app.setOrganizationName("WILOPRO")
    window = MainWindow()
    window.show()
    sys.exit(app.exec())
